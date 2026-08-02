from __future__ import annotations

import json
import math
import os
import re
import shutil
import textwrap
import unicodedata
import urllib.request
from collections import Counter
from pathlib import Path

import openpyxl
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[2]
WORKSPACE = PROJECT.parent
SOURCE_DATA = WORKSPACE / "reference-dashboard-municipal" / "src" / "data"
TEMP = Path(os.environ.get("TEMP", str(PROJECT / ".tmp"))) / "PlanProvincialMTS-20260801"
CHARTS = TEMP / "charts"
OUTPUT = (
    PROJECT
    / "public"
    / "downloads"
    / "planes-provinciales"
    / "03140000_Plan_Provincial_Maria_Trinidad_Sanchez_Documento_Base_2026.docx"
)

DASHBOARD_URL = "https://prodecare.net/DDPT/Dashboard-Territorial/data/territorial-dashboard.json"
DEMANDS_URL = "https://prodecare.net/DDPT/DemandasProvinciales/data/demandas_consolidadas_003.xlsx"
REFERENCE_URL = "https://prodecare.net/DDPT/planificacion-municipal/downloads/pmd-borradores/03140002_PMD_cabrera_Borrador_Tecnico_2025-2028.docx"

DASHBOARD_PATH = TEMP / "territorial-dashboard.json"
DEMANDS_PATH = TEMP / "demandas_consolidadas_003.xlsx"
REFERENCE_PATH = TEMP / "03140002_PMD_cabrera_Borrador_Tecnico_2025-2028.docx"

PROVINCE = "María Trinidad Sánchez"
PROVINCE_KEY = "mariatrinidadsanchez"
REGION = "Cibao Nordeste"

COLORS = {
    "ink": "#203740",
    "muted": "#65777E",
    "blue": "#2E74B5",
    "green": "#13836D",
    "dark_green": "#163F37",
    "gold": "#B57A16",
    "orange": "#D66A1D",
    "red": "#D95A66",
    "purple": "#7A5AF8",
    "cyan": "#34B3E5",
    "teal": "#4FA79D",
    "light": "#F4F7F7",
    "line": "#CFDADB",
    "white": "#FFFFFF",
    "pale_green": "#EEF7F4",
    "pale_gold": "#FCF6E9",
    "pale_blue": "#EEF5FA",
    "pale_red": "#FCF0F2",
}

W, H = 1300, 1560


def ensure_download(url: str, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists() or path.stat().st_size == 0:
        request = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(request, timeout=60) as response, path.open("wb") as target:
            shutil.copyfileobj(response, target)
    return path


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def normalize(value) -> str:
    text = unicodedata.normalize("NFD", str(value or ""))
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return re.sub(r"[^a-z0-9]", "", text.lower())


def find_record(path: Path, province: str = PROVINCE):
    data = load_json(path)
    found = []

    def walk(value):
        if isinstance(value, list):
            for item in value:
                walk(item)
        elif isinstance(value, dict):
            if value.get("provincia") == province:
                found.append(value)
            for item in value.values():
                walk(item)

    walk(data)
    return found[0] if found else None


def parse_demands(path: Path):
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheet = workbook["Demanda 2026"]
    headers = [str(value or "").strip() for value in next(sheet.iter_rows(min_row=4, max_row=4, values_only=True))]
    index = {normalize(value): position for position, value in enumerate(headers) if value}

    def get(values, aliases):
        for alias in aliases:
            key = normalize(alias)
            if key in index:
                value = values[index[key]]
                return "" if value is None else str(value).strip()
        return ""

    rows = []
    for values in sheet.iter_rows(min_row=5, values_only=True):
        province = get(values, ["Provincia"])
        if normalize(province) != normalize(PROVINCE):
            continue
        rows.append(
            {
                "number": get(values, ["No. X CDP", "No."]),
                "demand": get(values, ["Demanda Priorizada"]),
                "institution": get(values, ["INSTITUCION RESPONSABLE"]),
                "ally": get(values, ["INSTITUCION ALIADA"]),
                "municipality": get(values, ["Municipio"]),
                "place": get(values, ["Lugar"]),
                "theme": get(values, ["Tema_Comun"]),
                "functional": get(values, ["Clasificador_Funcional"]),
                "snip": get(values, ["Codigo SNIP"]),
                "year": get(values, ["año"]),
            }
        )
    return rows


def shorten_institution(value: str) -> str:
    value = re.sub(r"\s+", " ", value or "").strip()
    acronyms = re.findall(r"\(([A-ZÁÉÍÓÚÑ]{2,12})\)", value)
    if acronyms:
        return acronyms[-1]
    replacements = {
        "INSTITUTO NACIONAL DE AGUAS POTABLES Y ALCANTARILLADOS": "INAPA",
        "MINISTERIO DE OBRAS PUBLICAS Y COMUNICACIONES": "MOPC",
        "MINISTERIO DE OBRAS PÚBLICAS Y COMUNICACIONES": "MOPC",
        "INSTITUTO NACIONAL DE RECURSOS HIDRAULICOS": "INDRHI",
        "INSTITUTO NACIONAL DE RECURSOS HIDRÁULICOS": "INDRHI",
        "EMPRESA DE TRANSMISION ELECTRICA DOMINICANA": "ETED",
        "EMPRESA DE TRANSMISIÓN ELÉCTRICA DOMINICANA": "ETED",
        "CONSEJO NACIONAL DE DISCAPACIDAD": "CONADIS",
    }
    upper = value.upper()
    for key, short in replacements.items():
        if key in upper:
            return short
    return value[:28] + ("…" if len(value) > 28 else "")


_font_cache = {}


def font(size: int, bold: bool = False):
    key = (size, bold)
    if key in _font_cache:
        return _font_cache[key]
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            _font_cache[key] = ImageFont.truetype(str(candidate), size=size)
            return _font_cache[key]
    _font_cache[key] = ImageFont.load_default()
    return _font_cache[key]


def rgb(hex_color: str):
    value = hex_color.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def make_canvas():
    return Image.new("RGB", (W, H), rgb(COLORS["light"]))


def rounded(draw, box, fill="#FFFFFF", outline=None, radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill), outline=rgb(outline) if outline else None, width=width)


def wrap_text(draw, text: str, fnt, max_width: int):
    lines = []
    for paragraph in str(text or "").splitlines() or [""]:
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = current + " " + word
            if draw.textlength(candidate, font=fnt) <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_wrapped(draw, xy, text, fnt, fill, max_width, line_gap=5, max_lines=None, anchor=None):
    x, y = xy
    lines = wrap_text(draw, text, fnt, max_width)
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        last = lines[-1]
        while last and draw.textlength(last + "…", font=fnt) > max_width:
            last = last[:-1]
        lines[-1] = last.rstrip() + "…"
    line_height = fnt.getbbox("Ag")[3] - fnt.getbbox("Ag")[1]
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=rgb(fill), anchor=anchor)
        y += line_height + line_gap
    return y


def plate_header(draw, title: str, subtitle: str, source: str):
    draw.text((70, 52), title, font=font(38, True), fill=rgb(COLORS["ink"]))
    draw.text((70, 104), subtitle, font=font(22), fill=rgb(COLORS["muted"]))
    draw.text((1230, 72), source, font=font(16), fill=rgb(COLORS["muted"]), anchor="ra")


def plate_footer(draw, text: str):
    draw.text((70, 1515), text, font=font(15), fill=rgb(COLORS["muted"]))


def metric_card(draw, box, label: str, value: str, color: str, note: str = ""):
    rounded(draw, box, COLORS["white"], COLORS["line"], radius=14, width=2)
    x1, y1, x2, y2 = box
    draw.text((x1 + 22, y1 + 18), label, font=font(18, True), fill=rgb(color))
    draw.text((x1 + 22, y1 + 51), value, font=font(31, True), fill=rgb(COLORS["ink"]))
    if note:
        draw.text((x1 + 22, y2 - 27), note, font=font(14), fill=rgb(COLORS["muted"]))


def section_card(draw, box, title: str, tint: str = "#FFFFFF", title_color: str | None = None):
    rounded(draw, box, tint, COLORS["line"], radius=14, width=2)
    x1, y1, _, _ = box
    draw.text((x1 + 24, y1 + 22), title, font=font(22, True), fill=rgb(title_color or COLORS["blue"]))


def horizontal_bars(draw, box, labels, values, color=None, value_format=lambda v: f"{v:,.0f}", max_value=None):
    x1, y1, x2, y2 = box
    max_value = max_value or max(values or [1]) or 1
    bar_area_x = x1 + 190
    bar_area_w = x2 - bar_area_x - 75
    row_h = max(36, int((y2 - y1) / max(1, len(values))))
    palette = color if isinstance(color, list) else [color or COLORS["blue"]] * len(values)
    for i, (label, value) in enumerate(zip(labels, values)):
        cy = y1 + i * row_h + row_h // 2
        draw.text((x1, cy), str(label), font=font(16), fill=rgb(COLORS["ink"]), anchor="lm")
        draw.rounded_rectangle((bar_area_x, cy - 9, bar_area_x + bar_area_w, cy + 9), radius=9, fill=rgb("#E4EBEC"))
        width = max(2, int(bar_area_w * float(value) / max_value))
        draw.rounded_rectangle((bar_area_x, cy - 9, bar_area_x + width, cy + 9), radius=9, fill=rgb(palette[i % len(palette)]))
        draw.text((x2, cy), value_format(value), font=font(15, True), fill=rgb(COLORS["ink"]), anchor="rm")


def grouped_bars(draw, box, labels, series, colors, value_format=lambda v: f"{v:,.0f}"):
    x1, y1, x2, y2 = box
    all_values = [value for _, values in series for value in values]
    max_value = max(all_values or [1])
    chart_top = y1 + 25
    chart_bottom = y2 - 55
    chart_h = chart_bottom - chart_top
    group_w = (x2 - x1 - 70) / max(1, len(labels))
    bar_w = max(8, group_w * 0.62 / max(1, len(series)))
    for tick in range(5):
        value = max_value * tick / 4
        yy = chart_bottom - chart_h * tick / 4
        draw.line((x1 + 50, yy, x2, yy), fill=rgb("#DCE4E5"), width=1)
        draw.text((x1 + 42, yy), value_format(value), font=font(12), fill=rgb(COLORS["muted"]), anchor="rm")
    for i, label in enumerate(labels):
        center = x1 + 60 + group_w * (i + 0.5)
        for j, (_, values) in enumerate(series):
            value = values[i]
            height = chart_h * value / max_value if max_value else 0
            bx1 = center - (len(series) * bar_w) / 2 + j * bar_w
            draw.rectangle((bx1, chart_bottom - height, bx1 + bar_w - 2, chart_bottom), fill=rgb(colors[j]))
        draw.text((center, chart_bottom + 12), str(label), font=font(13), fill=rgb(COLORS["ink"]), anchor="ma")
    legend_x = x1 + 60
    for i, (name, _) in enumerate(series):
        draw.rectangle((legend_x, y1, legend_x + 18, y1 + 18), fill=rgb(colors[i]))
        draw.text((legend_x + 26, y1 + 9), name, font=font(14), fill=rgb(COLORS["ink"]), anchor="lm")
        legend_x += 190


def line_chart(draw, box, labels, values, color, value_format=lambda v: f"{v:.1f}", highlight_last=True):
    x1, y1, x2, y2 = box
    left, right, top, bottom = x1 + 55, x2 - 20, y1 + 20, y2 - 50
    minimum = min(values or [0])
    maximum = max(values or [1])
    span = maximum - minimum or 1
    floor = max(0, minimum - span * 0.15)
    ceiling = maximum + span * 0.15
    for tick in range(5):
        value = floor + (ceiling - floor) * tick / 4
        yy = bottom - (bottom - top) * tick / 4
        draw.line((left, yy, right, yy), fill=rgb("#DDE5E6"), width=1)
        draw.text((left - 10, yy), value_format(value), font=font(12), fill=rgb(COLORS["muted"]), anchor="rm")
    points = []
    for i, value in enumerate(values):
        xx = left + (right - left) * i / max(1, len(values) - 1)
        yy = bottom - (bottom - top) * (value - floor) / (ceiling - floor)
        points.append((xx, yy))
    if len(points) > 1:
        draw.line(points, fill=rgb(color), width=5, joint="curve")
    for i, point in enumerate(points):
        radius = 7 if highlight_last and i == len(points) - 1 else 4
        draw.ellipse((point[0] - radius, point[1] - radius, point[0] + radius, point[1] + radius), fill=rgb(color))
    label_step = max(1, math.ceil(len(labels) / 8))
    for i, label in enumerate(labels):
        if i % label_step == 0 or i == len(labels) - 1:
            xx = left + (right - left) * i / max(1, len(labels) - 1)
            draw.text((xx, bottom + 15), str(label), font=font(12), fill=rgb(COLORS["muted"]), anchor="ma")


def donut(draw, box, labels, values, colors, center_value="", center_label=""):
    x1, y1, x2, y2 = box
    diameter = min(y2 - y1 - 30, x2 - x1 - 260)
    diameter = max(150, diameter)
    cx = x1 + diameter / 2 + 10
    cy = (y1 + y2) / 2
    total = sum(values) or 1
    start = -90
    for value, color in zip(values, colors):
        end = start + 360 * value / total
        draw.pieslice((cx - diameter / 2, cy - diameter / 2, cx + diameter / 2, cy + diameter / 2), start=start, end=end, fill=rgb(color))
        start = end
    inner = diameter * 0.58
    draw.ellipse((cx - inner / 2, cy - inner / 2, cx + inner / 2, cy + inner / 2), fill=rgb(COLORS["white"]))
    if center_value:
        draw.text((cx, cy - 10), center_value, font=font(27, True), fill=rgb(COLORS["ink"]), anchor="mm")
    if center_label:
        draw.text((cx, cy + 28), center_label, font=font(14), fill=rgb(COLORS["muted"]), anchor="mm")
    lx = x1 + diameter + 45
    row_h = max(28, int((y2 - y1 - 20) / max(1, len(labels))))
    for i, (label, value, color) in enumerate(zip(labels, values, colors)):
        yy = y1 + 20 + i * row_h
        draw.rectangle((lx, yy + 2, lx + 17, yy + 19), fill=rgb(color))
        draw.text((lx + 28, yy), str(label), font=font(15), fill=rgb(COLORS["ink"]))
        draw.text((x2 - 10, yy), f"{value / total * 100:.1f}%", font=font(15, True), fill=rgb(COLORS["ink"]), anchor="ra")


def draw_geo_map(draw, box, geojson, highlight_key=PROVINCE_KEY):
    x1, y1, x2, y2 = box
    features = geojson.get("features", [])
    all_points = []

    def polygon_rings(geometry):
        if not geometry:
            return []
        coords = geometry.get("coordinates", [])
        if geometry.get("type") == "Polygon":
            return [coords]
        if geometry.get("type") == "MultiPolygon":
            return coords
        return []

    for feature in features:
        for polygon in polygon_rings(feature.get("geometry")):
            if polygon:
                all_points.extend(polygon[0])
    if not all_points:
        return
    min_x = min(p[0] for p in all_points)
    max_x = max(p[0] for p in all_points)
    min_y = min(p[1] for p in all_points)
    max_y = max(p[1] for p in all_points)
    scale = min((x2 - x1) / (max_x - min_x), (y2 - y1) / (max_y - min_y)) * 0.92
    ox = (x1 + x2) / 2 - (min_x + max_x) / 2 * scale
    oy = (y1 + y2) / 2 + (min_y + max_y) / 2 * scale

    def project(point):
        return (ox + point[0] * scale, oy - point[1] * scale)

    for feature in features:
        props = feature.get("properties", {})
        key = normalize(props.get("provinceKey") or props.get("shapeName") or props.get("name"))
        fill = COLORS["green"] if key == highlight_key else "#DDE5E5"
        for polygon in polygon_rings(feature.get("geometry")):
            if polygon and len(polygon[0]) >= 3:
                draw.polygon([project(point) for point in polygon[0]], fill=rgb(fill), outline=rgb(COLORS["white"]))


def save_plate(image, name):
    CHARTS.mkdir(parents=True, exist_ok=True)
    path = CHARTS / name
    image.save(path, "PNG", optimize=True)
    return path


def build_plates(facts):
    province = facts["province"]
    metrics = province["metrics"]
    municipalities = sorted(facts["municipalities"], key=lambda item: item["population"], reverse=True)
    municipal_condition = facts["municipal_condition"]
    urban = facts["urban_rural"]
    households = facts["households"]
    pyramid = facts["pyramid"]
    education_level = facts["education_level"]
    education = facts["education"]
    economy = facts["economy"]
    health = facts["health"]
    investment_2026 = facts["investment_2026"]
    demands = facts["demands"]
    geojson = facts["geojson"]
    paths = []

    # Lámina 1 — Territorio y distribución municipal
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Territorio y población", PROVINCE, "Censo 2022")
    cards = [(70 + i * 300, 145, 350 + i * 300, 270) for i in range(4)]
    metric_card(draw, cards[0], "Población", f"{province['population']:,}", COLORS["blue"], "habitantes")
    metric_card(draw, cards[1], "Municipios", str(province["municipalityCount"]), COLORS["green"], "Nagua, Cabrera, El Factor y Río San Juan")
    metric_card(draw, cards[2], "Zona urbana", f"{urban['urbana'] / urban['poblacion_total'] * 100:.1f}%", COLORS["purple"], f"{urban['urbana']:,} habitantes")
    metric_card(draw, cards[3], "Zona rural", f"{urban['rural'] / urban['poblacion_total'] * 100:.1f}%", COLORS["orange"], f"{urban['rural']:,} habitantes")
    section_card(draw, (70, 305, 650, 875), "Ubicación nacional")
    draw_geo_map(draw, (105, 385, 615, 790), geojson)
    draw_wrapped(draw, (105, 812), "La provincia se resalta únicamente como referencia cartográfica. La lectura intraprovincial requiere la localización posterior por municipio, distrito municipal y comunidad.", font(17), COLORS["muted"], 500, 5, 4)
    section_card(draw, (680, 305, 1230, 875), "Población por municipio")
    horizontal_bars(
        draw,
        (710, 385, 1195, 785),
        [item["name"] for item in municipalities],
        [item["population"] for item in municipalities],
        [COLORS["blue"], COLORS["green"], COLORS["teal"], COLORS["gold"]],
        value_format=lambda value: f"{value:,.0f}",
    )
    total = province["population"]
    y = 770
    for item in municipalities:
        draw.text((720, y), item["name"], font=font(14), fill=rgb(COLORS["ink"]))
        draw.text((1185, y), f"{item['population'] / total * 100:.1f}% del total", font=font(14, True), fill=rgb(COLORS["muted"]), anchor="ra")
        y += 22
    section_card(draw, (70, 910, 1230, 1465), "Preguntas para la validación territorial", COLORS["pale_blue"])
    questions = [
        "¿La distribución municipal de la población coincide con los patrones actuales de movilidad y uso de servicios?",
        "¿Qué comunidades quedan ocultas por el agregado provincial y deben localizarse antes de acordar prioridades?",
        "¿Qué datos de superficie, accesibilidad y riesgo debe incorporar la comisión técnica del CDP?",
    ]
    y = 980
    for number, question in enumerate(questions, 1):
        draw.ellipse((105, y, 139, y + 34), fill=rgb(COLORS["blue"]))
        draw.text((122, y + 17), str(number), font=font(16, True), fill=rgb(COLORS["white"]), anchor="mm")
        y = draw_wrapped(draw, (160, y + 2), question, font(20), COLORS["ink"], 1000, 6, 3) + 34
    plate_footer(draw, "Fuente: X Censo Nacional de Población y Vivienda 2022, ONE; cartografía ADM1 del portal provincial.")
    paths.append(save_plate(image, "lamina_01_territorio.png"))

    # Lámina 2 — Estructura demográfica y hogares
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Estructura demográfica", PROVINCE, "Censo 2022")
    metric_card(draw, (70, 145, 370, 270), "Hogares", f"{households['hogares_total']:,}", COLORS["blue"], "hogares censados")
    metric_card(draw, (395, 145, 695, 270), "Población en hogares", f"{households['poblacion_en_hogares']:,}", COLORS["green"], "personas")
    metric_card(draw, (720, 145, 1020, 270), "Personas por hogar", f"{households['personas_por_hogar']:.2f}", COLORS["purple"], "promedio")
    all_age = pyramid["age_groups"]
    male_total = sum(item["male"] for item in all_age)
    female_total = sum(item["female"] for item in all_age)
    metric_card(draw, (1045, 145, 1230, 270), "H / M", f"{male_total / (male_total + female_total) * 100:.1f} / {female_total / (male_total + female_total) * 100:.1f}", COLORS["orange"], "%")
    section_card(draw, (70, 305, 1230, 1160), "Pirámide de población 2022")
    age_groups = [item for item in all_age if item["age_group"] != "No declarado"]
    age_groups = list(reversed(age_groups))
    max_age = max(max(item["male"], item["female"]) for item in age_groups)
    center = 650
    top, bottom = 385, 1085
    row_h = (bottom - top) / len(age_groups)
    half_width = 480
    for i, item in enumerate(age_groups):
        cy = top + (i + 0.5) * row_h
        mw = half_width * item["male"] / max_age
        fw = half_width * item["female"] / max_age
        draw.rectangle((center - mw, cy - row_h * 0.36, center, cy + row_h * 0.36), fill=rgb(COLORS["blue"]))
        draw.rectangle((center, cy - row_h * 0.36, center + fw, cy + row_h * 0.36), fill=rgb(COLORS["red"]))
        if i % 2 == 0:
            draw.text((center, cy), item["age_group"].strip(), font=font(12), fill=rgb(COLORS["ink"]), anchor="mm")
    draw.line((center, top - 10, center, bottom + 10), fill=rgb(COLORS["white"]), width=2)
    draw.text((255, 1120), "Hombres", font=font(18, True), fill=rgb(COLORS["blue"]), anchor="mm")
    draw.text((1045, 1120), "Mujeres", font=font(18, True), fill=rgb(COLORS["red"]), anchor="mm")
    age_0_14 = sum(item["male"] + item["female"] for item in all_age[:4])
    age_65 = sum(item["male"] + item["female"] for item in all_age[15:] if item["age_group"] != "No declarado")
    population_pyramid = male_total + female_total
    age_15_64 = population_pyramid - age_0_14 - age_65
    section_card(draw, (70, 1195, 1230, 1465), "Composición por grandes grupos", COLORS["pale_green"])
    horizontal_bars(
        draw,
        (105, 1265, 1190, 1430),
        ["0–14 años", "15–64 años", "65 años o más"],
        [age_0_14, age_15_64, age_65],
        [COLORS["cyan"], COLORS["green"], COLORS["gold"]],
        value_format=lambda value: f"{value / population_pyramid * 100:.1f}%",
        max_value=max(age_0_14, age_15_64, age_65),
    )
    plate_footer(draw, "Fuente: X Censo Nacional de Población y Vivienda 2022, ONE. Los grupos deben cruzarse con sexo y localización antes de decidir.")
    paths.append(save_plate(image, "lamina_02_demografia.png"))

    # Lámina 3 — Servicios y condición de vida
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Condición de vida y servicios", PROVINCE, "Hogares · Censo 2022")
    services = facts["condition"]["servicios"]
    household_total = services["servicios_sanitarios"]["total"]
    water_pct = services["agua_uso_domestico"]["categorias"]["del_acueducto_dentro_de_la_vivienda"] / household_total * 100
    toilet_pct = services["servicios_sanitarios"]["categorias"]["inodoro"] / household_total * 100
    garbage_pct = services["eliminacion_basura"]["categorias"]["la_recoge_el_ayuntamiento"] / household_total * 100
    internet_pct = facts["tic"]["internet"]["rate_used"] * 100
    cards = [(70 + i * 300, 145, 350 + i * 300, 270) for i in range(4)]
    metric_card(draw, cards[0], "Acueducto dentro", f"{water_pct:.1f}%", COLORS["blue"], "hogares")
    metric_card(draw, cards[1], "Hogares con inodoro", f"{toilet_pct:.1f}%", COLORS["green"], "hogares")
    metric_card(draw, cards[2], "Recogida municipal", f"{garbage_pct:.1f}%", COLORS["orange"], "hogares")
    metric_card(draw, cards[3], "Uso de internet", f"{internet_pct:.1f}%", COLORS["purple"], "personas de referencia")
    section_card(draw, (70, 305, 1230, 930), "Cobertura declarada por municipio")
    names = []
    water_values, toilet_values, garbage_values = [], [], []
    condition_by_name = {item["municipio"]: item for item in municipal_condition}
    for item in municipalities:
        name = item["name"]
        record = condition_by_name[name]["servicios"]
        total_h = record["servicios_sanitarios"]["total"]
        names.append(name)
        water_values.append(record["agua_uso_domestico"]["categorias"]["del_acueducto_dentro_de_la_vivienda"] / total_h * 100)
        toilet_values.append(record["servicios_sanitarios"]["categorias"]["inodoro"] / total_h * 100)
        garbage_values.append(record["eliminacion_basura"]["categorias"]["la_recoge_el_ayuntamiento"] / total_h * 100)
    left = 105
    bar_area_x = 320
    chart_right = 1175
    row_space = 130
    for i, name in enumerate(names):
        y = 400 + i * row_space
        draw.text((left, y + 25), name, font=font(18, True), fill=rgb(COLORS["ink"]), anchor="lm")
        for j, (label, value, color) in enumerate(
            [("Agua", water_values[i], COLORS["blue"]), ("Inodoro", toilet_values[i], COLORS["green"]), ("Residuos", garbage_values[i], COLORS["orange"])]
        ):
            yy = y + j * 31
            draw.text((235, yy), label, font=font(14), fill=rgb(COLORS["muted"]), anchor="rm")
            draw.rounded_rectangle((bar_area_x, yy - 7, chart_right, yy + 7), radius=7, fill=rgb("#E4EBEC"))
            draw.rounded_rectangle((bar_area_x, yy - 7, bar_area_x + (chart_right - bar_area_x) * value / 100, yy + 7), radius=7, fill=rgb(color))
            draw.text((1195, yy), f"{value:.1f}%", font=font(14, True), fill=rgb(COLORS["ink"]), anchor="rm")
    garbage = services["eliminacion_basura"]["categorias"]
    section_card(draw, (70, 965, 650, 1465), "Eliminación de residuos declarada", COLORS["pale_gold"], COLORS["orange"])
    donut(
        draw,
        (100, 1040, 625, 1425),
        ["Recogida ayuntamiento", "Quema", "Otros"],
        [garbage["la_recoge_el_ayuntamiento"], garbage["la_queman"], household_total - garbage["la_recoge_el_ayuntamiento"] - garbage["la_queman"]],
        [COLORS["green"], COLORS["orange"], COLORS["muted"]],
        center_value=f"{household_total:,}",
        center_label="hogares",
    )
    section_card(draw, (680, 965, 1230, 1465), "Preguntas para el CDP", COLORS["pale_blue"])
    questions = [
        "¿En qué comunidades se concentran las diferencias entre municipios?",
        "¿La cobertura declarada coincide con continuidad, calidad, frecuencia y costo del servicio?",
        "¿Qué competencias corresponden a gobiernos locales y cuáles requieren coordinación sectorial?",
    ]
    y = 1045
    for question in questions:
        draw.ellipse((715, y + 5, 729, y + 19), fill=rgb(COLORS["blue"]))
        y = draw_wrapped(draw, (750, y), question, font(18), COLORS["ink"], 430, 5, 4) + 22
    plate_footer(draw, "Fuente: X Censo 2022, ONE, según datasets del Dashboard de Diagnóstico Territorial. Cobertura declarada no equivale a calidad.")
    paths.append(save_plate(image, "lamina_03_servicios.png"))

    # Lámina 4 — Condiciones sociales
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Condiciones sociales", PROVINCE, "Registros 2026")
    overcrowding = metrics["overcrowding"]
    disability = metrics["disability"]
    inaipi = metrics["inaipi"]
    metric_card(draw, (70, 145, 350, 270), "Hacinamiento extremo", f"{overcrowding['extremePct']:.1f}%", COLORS["orange"], f"{overcrowding['extreme']:,} hogares")
    metric_card(draw, (370, 145, 650, 270), "Hacinamiento moderado", f"{overcrowding['moderatePct']:.1f}%", COLORS["gold"], f"{overcrowding['moderate']:,} hogares")
    metric_card(draw, (670, 145, 950, 270), "Centros INAIPI", f"{inaipi['centers']:,}", COLORS["green"], "registro disponible")
    metric_card(draw, (970, 145, 1230, 270), "Asistencia INAIPI", f"{inaipi['attendancePct']:.1f}%", COLORS["blue"], "presencias / registros")
    section_card(draw, (70, 305, 650, 865), "Distribución de hacinamiento", COLORS["white"])
    donut(
        draw,
        (100, 390, 625, 820),
        ["Extremo", "Moderado", "Leve", "Sin hacinamiento"],
        [overcrowding["extreme"], overcrowding["moderate"], overcrowding["mild"], overcrowding["notOvercrowded"]],
        [COLORS["red"], COLORS["orange"], COLORS["gold"], COLORS["green"]],
        center_value=f"{overcrowding['total']:,}",
        center_label="hogares",
    )
    section_card(draw, (680, 305, 1230, 865), "Personas con discapacidad por ICV", COLORS["white"])
    horizontal_bars(
        draw,
        (715, 400, 1195, 790),
        ["ICV 1", "ICV 2", "ICV 3", "ICV 4"],
        [disability["icv1"], disability["icv2"], disability["icv3"], disability["icv4"]],
        [COLORS["red"], COLORS["orange"], COLORS["gold"], COLORS["green"]],
        value_format=lambda value: f"{value:,.0f}",
    )
    section_card(draw, (70, 900, 1230, 1195), "Registro de atención INAIPI", COLORS["pale_green"], COLORS["green"])
    registered = inaipi["registered"]
    present = inaipi["present"]
    draw.text((110, 990), "Registrados", font=font(19), fill=rgb(COLORS["ink"]))
    draw.rounded_rectangle((300, 984, 1165, 1012), radius=14, fill=rgb("#DCE7E4"))
    draw.rounded_rectangle((300, 984, 1165, 1012), radius=14, fill=rgb(COLORS["muted"]))
    draw.text((1190, 998), f"{registered:,}", font=font(18, True), fill=rgb(COLORS["ink"]), anchor="rm")
    draw.text((110, 1065), "Presentes", font=font(19), fill=rgb(COLORS["ink"]))
    draw.rounded_rectangle((300, 1059, 1165, 1087), radius=14, fill=rgb("#DCE7E4"))
    draw.rounded_rectangle((300, 1059, 300 + 865 * present / registered, 1087), radius=14, fill=rgb(COLORS["green"]))
    draw.text((1190, 1073), f"{present:,}", font=font(18, True), fill=rgb(COLORS["ink"]), anchor="rm")
    draw_wrapped(draw, (110, 1125), "La fuente registra atenciones o presencias del período; no debe interpretarse automáticamente como número de personas únicas ni como cobertura efectiva de toda la población objetivo.", font(16), COLORS["muted"], 1080, 4, 3)
    section_card(draw, (70, 1230, 1230, 1465), "Preguntas para el CDP", COLORS["pale_blue"])
    draw_wrapped(draw, (105, 1305), "¿Dónde se localizan los hogares con mayor hacinamiento? ¿Qué población con discapacidad enfrenta barreras de acceso? ¿Qué información nominal, territorial y de capacidad institucional debe validarse antes de formular resultados?", font(20), COLORS["ink"], 1085, 8, 5)
    plate_footer(draw, "Fuente: Dashboard Territorial; cortes administrativos 2026-2. Las definiciones ICV y de asistencia deben revisarse en su ficha técnica.")
    paths.append(save_plate(image, "lamina_04_social.png"))

    # Lámina 5 — Seguridad y movilidad
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Seguridad vial y convivencia", PROVINCE, "Series disponibles")
    homicide = [item for item in metrics["homicide"]["series"] if item["year"] <= 2024]
    traffic = [item for item in metrics["traffic"]["series"] if item["year"] <= 2025]
    metric_card(draw, (70, 145, 350, 270), "Homicidios", f"{metrics['homicide']['latest']['count']}", COLORS["red"], "2024")
    metric_card(draw, (370, 145, 650, 270), "Tasa de homicidios", f"{metrics['homicide']['latest']['rate']:.1f}", COLORS["red"], "por 100 mil · 2024")
    metric_card(draw, (670, 145, 950, 270), "Muertes de tránsito", f"{traffic[-1]['count']}", COLORS["orange"], "2025")
    metric_card(draw, (970, 145, 1230, 270), "Tasa de tránsito", f"{traffic[-1]['rate']:.1f}", COLORS["orange"], "por 100 mil · 2025")
    section_card(draw, (70, 305, 1230, 760), "Tasa de homicidios por 100 mil habitantes")
    line_chart(draw, (105, 385, 1195, 720), [item["year"] for item in homicide], [item["rate"] for item in homicide], COLORS["red"])
    section_card(draw, (70, 795, 1230, 1235), "Tasa de muertes por tránsito por 100 mil habitantes")
    line_chart(draw, (105, 875, 1195, 1190), [item["year"] for item in traffic], [item["rate"] for item in traffic], COLORS["orange"])
    roads = metrics["roads"]["latest"]
    section_card(draw, (70, 1270, 1230, 1465), "Dato complementario y preguntas", COLORS["pale_blue"])
    draw_wrapped(draw, (105, 1335), f"La fuente registra {roads['records']} intervenciones viales, {roads['lengthKm']:.1f} km y {roads['areaM2']:,.0f} m² en el período {roads['period']}. ¿Dónde se concentran los eventos? ¿Qué vías, horarios y grupos requieren verificación local? ¿Qué parte del problema corresponde a infraestructura, control, educación o atención de emergencias?", font(17), COLORS["ink"], 1085, 5, 5)
    plate_footer(draw, "Fuente: Dashboard Territorial. Se excluye 2026 de las series comparativas por corresponder a un período parcial.")
    paths.append(save_plate(image, "lamina_05_seguridad.png"))

    # Lámina 6 — Educación y economía
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Educación y economía formal", PROVINCE, "Censo 2022 · Anuario/DEE 2024")
    levels = education_level["nivel"]
    level_labels = ["Ninguno", "Preprimaria", "Primaria", "Secundaria", "Superior"]
    level_values = [levels[key]["total"] for key in ["ninguno", "preprimaria", "primaria", "secundaria", "superior"]]
    level_total = sum(level_values)
    efficiency = education["anuario"]["eficiencia"]["secundario"]
    dee = economy["dee_2024"]
    metric_card(draw, (70, 145, 350, 270), "Establecimientos", f"{dee['total_establishments']:,}", COLORS["red"], "DEE 2024")
    metric_card(draw, (370, 145, 650, 270), "Empleo estimado", f"{dee['total_employees']:,.0f}", COLORS["red"], "DEE 2024")
    metric_card(draw, (670, 145, 950, 270), "Promoción secundaria", f"{efficiency['promocion']:.1f}%", COLORS["green"], "Anuario 2024")
    metric_card(draw, (970, 145, 1230, 270), "Abandono secundario", f"{efficiency['abandono']:.1f}%", COLORS["orange"], "Anuario 2024")
    section_card(draw, (70, 305, 650, 865), "Nivel de instrucción · distribución", COLORS["pale_blue"])
    horizontal_bars(
        draw,
        (105, 390, 615, 800),
        level_labels,
        level_values,
        [COLORS["purple"], COLORS["orange"], COLORS["cyan"], COLORS["green"], COLORS["blue"]],
        value_format=lambda value: f"{value / level_total * 100:.1f}%",
    )
    draw_wrapped(draw, (105, 805), "Distribución según las categorías consolidadas del Dashboard; usar como composición relativa y validar denominador antes de estimar población objetivo.", font(14), COLORS["muted"], 500, 4, 3)
    section_card(draw, (680, 305, 1230, 865), "Establecimientos por tamaño", COLORS["pale_red"], COLORS["red"])
    bands = dee["employment_size_bands"]
    horizontal_bars(
        draw,
        (715, 390, 1195, 800),
        ["Micro", "Pequeña", "Mediana", "Grande"],
        [item["establishments"] for item in bands],
        [COLORS["cyan"], COLORS["green"], COLORS["gold"], COLORS["red"]],
        value_format=lambda value: f"{value:,.0f}",
    )
    section_card(draw, (70, 900, 1230, 1295), "Secciones CIIU con mayor empleo estimado", COLORS["white"])
    top_sectors = sorted(dee["sectors"], key=lambda item: item["employees"], reverse=True)[:5]
    horizontal_bars(
        draw,
        (105, 985, 1195, 1245),
        [item["label"].split(";")[0][:29] for item in top_sectors],
        [item["employees"] for item in top_sectors],
        COLORS["red"],
        value_format=lambda value: f"{value:,.0f}",
    )
    section_card(draw, (70, 1330, 1230, 1465), "Pregunta para el CDP", COLORS["pale_green"])
    draw_wrapped(draw, (105, 1380), "¿Qué brechas educativas y capacidades económicas son verificables por municipio, sexo, edad y zona? ¿Qué parte de la economía informal o agropecuaria no está representada en el DEE?", font(18), COLORS["ink"], 1080, 5, 3)
    plate_footer(draw, "Fuente: X Censo 2022; Anuario Estadístico Educativo 2024; Directorio de Empresas y Establecimientos 2024.")
    paths.append(save_plate(image, "lamina_06_educacion_economia.png"))

    # Lámina 7 — Salud y deporte
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Equipamientos de salud y deporte", PROVINCE, "Registros disponibles")
    health_counts = Counter(item["tipo_centro"].title() for item in health["centros"])
    sports = metrics["sports"]
    metric_card(draw, (70, 145, 350, 270), "Centros de salud", f"{len(health['centros']):,}", COLORS["gold"], "registros SNS disponibles")
    metric_card(draw, (370, 145, 650, 270), "Instalaciones deportivas", f"{sports['count']:,}", COLORS["green"], "2025")
    metric_card(draw, (670, 145, 950, 270), "Por 10 mil hab.", f"{sports['per10k']:.1f}", COLORS["green"], "instalaciones deportivas")
    metric_card(draw, (970, 145, 1230, 270), "Propiedad municipal", f"{sports['owners'].get('El Ayuntamiento', 0)}", COLORS["blue"], "instalaciones")
    section_card(draw, (70, 305, 650, 1055), "Establecimientos de salud por tipo", COLORS["pale_gold"], COLORS["gold"])
    health_sorted = health_counts.most_common()
    horizontal_bars(
        draw,
        (105, 390, 615, 990),
        [label.replace("Centro De", "Centro de")[:25] for label, _ in health_sorted],
        [value for _, value in health_sorted],
        COLORS["gold"],
        value_format=lambda value: f"{value:,.0f}",
    )
    section_card(draw, (680, 305, 1230, 1055), "Instalaciones deportivas por propietario", COLORS["pale_green"], COLORS["green"])
    owner_counts = Counter(sports["owners"])
    owners = owner_counts.most_common()
    owner_labels = [
        "Ayuntamientos" if label == "El Ayuntamiento" else "MIDEREC" if label == "MIDEREC" else label.replace("Entidad ", "")[:23]
        for label, _ in owners
    ]
    horizontal_bars(
        draw,
        (715, 390, 1195, 990),
        owner_labels,
        [value for _, value in owners],
        [COLORS["green"], COLORS["blue"], COLORS["teal"], COLORS["gold"], COLORS["purple"], COLORS["orange"], COLORS["muted"], COLORS["cyan"]],
        value_format=lambda value: f"{value:,.0f}",
    )
    section_card(draw, (70, 1090, 1230, 1465), "Lectura y preguntas", COLORS["pale_blue"])
    draw_wrapped(draw, (105, 1165), "Los conteos describen presencia física registrada; no informan cartera de servicios, personal, estado, horario, accesibilidad, capacidad ni uso. ¿Qué equipamientos sirven a más de un municipio? ¿Dónde existen barreras de distancia o capacidad? ¿Qué inventario debe actualizarse antes de acordar resultados?", font(20), COLORS["ink"], 1080, 7, 7)
    plate_footer(draw, "Fuente: registros SNS disponibles en el Dashboard de Diagnóstico Territorial; inventario deportivo 2025 del Dashboard Territorial.")
    paths.append(save_plate(image, "lamina_07_salud_deporte.png"))

    # Lámina 8 — Inversión pública
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Inversión pública", PROVINCE, "Series y corte 2026")
    inv_series = metrics["investment"]["series"]
    metric_card(draw, (70, 145, 350, 270), "Proyectos 2026", f"{investment_2026['projectCount']:,}", COLORS["blue"], "asociados a la provincia")
    metric_card(draw, (370, 145, 650, 270), "Presupuesto 2026", f"RD$ {investment_2026['budget']/1e9:.2f} mil M", COLORS["green"], "corte 31-07-2026")
    metric_card(draw, (670, 145, 950, 270), "Ejecutado 2026", f"RD$ {investment_2026['executed']/1e9:.2f} mil M", COLORS["green"], "corte 31-07-2026")
    metric_card(draw, (970, 145, 1230, 270), "Ejecución 2026", f"{investment_2026['executionPct']:.1f}%", COLORS["orange"], "23 proyectos con ejecución")
    section_card(draw, (70, 305, 1230, 1045), "Presupuesto y ejecución registrada · 2018–2025")
    grouped_bars(
        draw,
        (105, 390, 1195, 985),
        [item["year"] for item in inv_series],
        [("Presupuesto", [item["budget"] / 1e9 for item in inv_series]), ("Ejecutado", [item["executed"] / 1e9 for item in inv_series])],
        [COLORS["blue"], COLORS["green"]],
        value_format=lambda value: f"{value:.1f}",
    )
    draw.text((115, 1002), "Escala: miles de millones de RD$", font=font(14), fill=rgb(COLORS["muted"]))
    section_card(draw, (70, 1080, 1230, 1275), "Lectura del corte 2026", COLORS["pale_green"], COLORS["green"])
    draw.text((110, 1148), "Ejecución registrada", font=font(18), fill=rgb(COLORS["ink"]))
    draw.rounded_rectangle((345, 1145, 1165, 1175), radius=15, fill=rgb("#DCE7E4"))
    draw.rounded_rectangle((345, 1145, 345 + 820 * investment_2026["executionPct"] / 100, 1175), radius=15, fill=rgb(COLORS["green"]))
    draw.text((1190, 1160), f"{investment_2026['executionPct']:.1f}%", font=font(18, True), fill=rgb(COLORS["ink"]), anchor="rm")
    draw_wrapped(draw, (110, 1205), f"Sector más frecuente por número de proyectos: {investment_2026['topSector']}. Hay {investment_2026['projectsWithActiveContracts']} proyectos con contratos activos registrados.", font(16), COLORS["muted"], 1050, 4, 3)
    section_card(draw, (70, 1310, 1230, 1465), "Pregunta para el CDP", COLORS["pale_blue"])
    draw_wrapped(draw, (105, 1360), "¿Qué inversiones responden a brechas verificadas, cuáles se concentran territorialmente y cuáles requieren coordinación, secuenciación o revisión de ejecución? La serie histórica y el corte 2026 no deben compararse sin revisar metodología y fecha.", font(17), COLORS["ink"], 1080, 5, 4)
    plate_footer(draw, "Fuente: Dashboard Territorial (serie 2018–2025) e Inversión Pública Territorial, corte 31-07-2026.")
    paths.append(save_plate(image, "lamina_08_inversion.png"))

    # Lámina 9 — Construcción y demandas
    image = make_canvas()
    draw = ImageDraw.Draw(image)
    plate_header(draw, "Diagnóstico provincial · Dinámica constructiva y demandas", PROVINCE, "Registros 2022–2026")
    permits = [item for item in metrics["permits"]["series"] if item["year"] <= 2025]
    latest_permit = metrics["permits"]["latest"]
    metric_card(draw, (70, 145, 350, 270), "Demandas CDP", f"{len(demands)}", COLORS["blue"], "consolidado 2026")
    metric_card(draw, (370, 145, 650, 270), "Licencias 2026", f"{latest_permit['licenses']}", COLORS["purple"], "enero–junio")
    metric_card(draw, (670, 145, 950, 270), "Área autorizada", f"{latest_permit['areaM2']:,.0f} m²", COLORS["purple"], "enero–junio 2026")
    metric_card(draw, (970, 145, 1230, 270), "Inversión declarada", f"RD$ {latest_permit['investment']/1e6:.1f} M", COLORS["purple"], "enero–junio 2026")
    section_card(draw, (70, 305, 650, 880), "Licencias e inversión declarada", COLORS["white"])
    years = [item["year"] for item in permits]
    investment_values = [item["investment"] / 1e6 for item in permits]
    line_chart(draw, (105, 390, 615, 760), years, investment_values, COLORS["purple"], value_format=lambda value: f"{value:.0f}")
    draw.text((110, 770), "Inversión declarada · millones de RD$", font=font(14), fill=rgb(COLORS["muted"]))
    y = 800
    for item in permits:
        draw.text((115, y), str(item["year"]), font=font(13, True), fill=rgb(COLORS["ink"]))
        draw.text((250, y), f"{item['licenses']} licencias", font=font(13), fill=rgb(COLORS["muted"]))
        draw.text((600, y), f"{item['areaM2']:,.0f} m²", font=font(13, True), fill=rgb(COLORS["ink"]), anchor="ra")
        y += 24
    section_card(draw, (680, 305, 1230, 880), "Demandas por tema común", COLORS["pale_blue"])
    theme_counts = Counter((item["theme"].split("-")[0] + "-" + item["theme"].split("-", 1)[1][:26]) if "-" in item["theme"] else item["theme"] for item in demands)
    horizontal_bars(
        draw,
        (715, 390, 1195, 820),
        [label for label, _ in theme_counts.most_common()],
        [value for _, value in theme_counts.most_common()],
        [COLORS["blue"], COLORS["green"], COLORS["orange"], COLORS["purple"], COLORS["gold"]],
        value_format=lambda value: f"{value:,.0f}",
    )
    section_card(draw, (70, 915, 650, 1260), "Localización declarada de las demandas", COLORS["pale_green"])
    municipality_counts = Counter(item["municipality"].title() if item["municipality"] else "Provincial / sin municipio" for item in demands)
    horizontal_bars(
        draw,
        (105, 1000, 615, 1205),
        [label for label, _ in municipality_counts.most_common()],
        [value for _, value in municipality_counts.most_common()],
        [COLORS["green"], COLORS["blue"], COLORS["teal"], COLORS["gold"]],
        value_format=lambda value: f"{value:,.0f}",
    )
    section_card(draw, (680, 915, 1230, 1260), "Instituciones responsables registradas", COLORS["pale_gold"], COLORS["orange"])
    institution_counts = Counter(shorten_institution(item["institution"]) for item in demands)
    horizontal_bars(
        draw,
        (715, 1000, 1195, 1205),
        [label for label, _ in institution_counts.most_common()],
        [value for _, value in institution_counts.most_common()],
        COLORS["orange"],
        value_format=lambda value: f"{value:,.0f}",
    )
    section_card(draw, (70, 1295, 1230, 1465), "Uso correcto en la formulación", COLORS["pale_blue"])
    draw_wrapped(draw, (105, 1348), "Las diez demandas son insumos consolidados del CDP. Su inclusión no las convierte automáticamente en acciones del plan ni acredita factibilidad, presupuesto, madurez técnica o aprobación. El CDP debe revisar su vigencia, alcance territorial, relación con el diagnóstico y estado SNIP.", font(17), COLORS["ink"], 1080, 5, 5)
    plate_footer(draw, "Fuente: permisos del Dashboard Territorial; Demandas Provinciales, libro consolidado 003, 2026.")
    paths.append(save_plate(image, "lamina_09_demandas.png"))

    return paths


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(inches)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths, header=True, font_size=9, alternating=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        if row_index == 0 and header:
            set_repeat_table_header(row)
        for column_index, cell in enumerate(row.cells):
            if column_index < len(widths):
                set_cell_width(cell, widths[column_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0 and header:
                set_cell_shading(cell, COLORS["ink"])
            elif alternating and row_index % 2 == 0:
                set_cell_shading(cell, "#F3F6F6")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor(255, 255, 255) if row_index == 0 and header else RGBColor.from_string("203740")
                    if row_index == 0 and header:
                        run.font.bold = True


def add_table(doc, headers, rows, widths, font_size=9, alternating=True):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
    style_table(table, widths, True, font_size, alternating)
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text_node, end])


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Calibri")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    run_properties.extend([fonts, color, underline, size])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def clear_container(container):
    element = container._element
    for child in list(element):
        element.remove(child)


def clear_document_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("203740")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 11, 5),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_header_footer(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header
    clear_container(header)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    set_cell_width(table.cell(0, 0), 3.25)
    set_cell_width(table.cell(0, 1), 3.25)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("PLAN PROVINCIAL DE DESARROLLO · DOCUMENTO BASE")
    set_run_font(run, size=7.5, color="13836D", bold=True)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    run = right.add_run("María Trinidad Sánchez · 03140000")
    set_run_font(run, size=7.5, color="65777E")
    footer = section.footer
    clear_container(footer)
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("María Trinidad Sánchez · Documento base para formulación · ")
    set_run_font(run, size=7.5, color="65777E")
    add_page_field(paragraph)


def add_body_paragraph(doc, text, bold_lead=None, small=False, source=False, center=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if source:
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(4)
    if bold_lead and text.startswith(bold_lead):
        run = paragraph.add_run(bold_lead)
        set_run_font(run, size=8.5 if source else (9 if small else 10.5), color="65777E" if source else "203740", bold=True)
        rest = text[len(bold_lead) :]
        run = paragraph.add_run(rest)
        set_run_font(run, size=8.5 if source else (9 if small else 10.5), color="65777E" if source else "203740")
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=8.5 if source else (9 if small else 10.5), color="65777E" if source else "203740")
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(3)
    for run in paragraph.runs:
        set_run_font(run, size=10, color="203740")
    if not paragraph.runs:
        run = paragraph.add_run(text)
        set_run_font(run, size=10, color="203740")
    else:
        paragraph.runs[0].text = text
    return paragraph


def add_note_box(doc, title, text, fill="EEF7F4", title_color="13836D"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_margins(cell, 120, 140, 120, 140)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=title_color, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color="203740")
    return table


def page_break(doc):
    doc.add_page_break()


def build_document(facts, plate_paths):
    doc = Document(str(REFERENCE_PATH))
    clear_document_body(doc)
    configure_styles(doc)
    configure_header_footer(doc)
    props = doc.core_properties
    props.title = "Documento base para la formulación del Plan Provincial de María Trinidad Sánchez"
    props.subject = "Base técnica para deliberación del Consejo de Desarrollo Provincial"
    props.author = "DDPT · Documento base generado para revisión institucional"
    props.keywords = "María Trinidad Sánchez; Plan Provincial; CDP; diagnóstico; inversión; demandas"

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REPÚBLICA DOMINICANA · PLANIFICACIÓN TERRITORIAL")
    set_run_font(r, size=10, color="B57A16", bold=True)
    p.paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento base para la formulación del")
    set_run_font(r, size=17, color="203740", bold=True)
    p.paragraph_format.space_after = Pt(3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Plan Estratégico de Desarrollo Territorial Provincial")
    set_run_font(r, size=24, color="203740", bold=True)
    p.paragraph_format.space_after = Pt(22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROVINCE)
    set_run_font(r, size=22, color="13836D", bold=True)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Región {REGION}")
    set_run_font(r, size=13, color="65777E")
    p.paragraph_format.space_after = Pt(26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BORRADOR TÉCNICO PARA REVISIÓN DEL CDP")
    set_run_font(r, size=13, color="B57A16", bold=True)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Diagnóstico, orientaciones de análisis y espacios para la formulación participativa")
    set_run_font(r, size=10, color="65777E")
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Código territorial 03140000 · Base estadística 2022–2026")
    set_run_font(r, size=9, color="65777E")
    page_break(doc)

    # Contents and synthesis
    doc.add_heading("Contenido", level=1)
    contents = [
        ("1", "Información general, marco jurídico e institucional y metodología"),
        ("2", "Diagnóstico provincial: nueve láminas de evidencia factual"),
        ("3", "Lectura técnica y matriz de fortalezas/debilidades para validación"),
        ("4", "Agenda de concertación y registro de acuerdos del CDP"),
        ("5", "Visión, objetivos, resultados y plan de acción por acordar"),
        ("Anexo", "Demandas provinciales consolidadas y fuentes de trazabilidad"),
    ]
    add_table(doc, ["Parte", "Contenido"], contents, [0.65, 5.85], font_size=9)
    doc.add_heading("Síntesis ejecutiva", level=1)
    add_body_paragraph(
        doc,
        "María Trinidad Sánchez registró 156,633 habitantes en el Censo 2022, distribuidos entre Nagua, Cabrera, El Factor y Río San Juan. Este documento organiza información factual sobre población, servicios, condiciones sociales, seguridad, educación, economía, salud, deporte, inversión pública, permisos y demandas provinciales.",
    )
    add_body_paragraph(
        doc,
        "El diagnóstico no sustituye la deliberación territorial. Las fortalezas y debilidades se presentan como hipótesis sustentadas en evidencia disponible; el Consejo de Desarrollo Provincial debe verificar su localización, vigencia, causalidad, competencia institucional y relevancia antes de convertirlas en visión, objetivos, resultados o acciones.",
    )
    add_note_box(
        doc,
        "Regla de uso",
        "Los datos y las diez demandas consolidadas están preelaborados como insumos. Ninguna visión, objetivo, acción, proyecto, meta, presupuesto, responsable o plazo ha sido decidido en este documento.",
        fill="FCF6E9",
        title_color="B57A16",
    )
    page_break(doc)

    # General information
    doc.add_heading("1. Información general", level=1)
    doc.add_heading("1.1 Equipo de formulación y participación del CDP", level=2)
    add_body_paragraph(
        doc,
        "La formulación se organiza por funciones institucionales y no por nombres de personas. El CDP deberá completar la integración efectiva, dejar constancia en actas y definir su mecanismo técnico de trabajo.",
    )
    roles = [
        ("Consejo de Desarrollo Provincial (CDP)", "Deliberar, concertar y validar el diagnóstico, la estrategia territorial, las prioridades y el seguimiento."),
        ("Coordinación del CDP", "Organizar convocatorias, agenda, articulación institucional y presentación de acuerdos."),
        ("Secretaría del CDP", "Registrar asistencia, actas, acuerdos, documentación y comunicación de las sesiones."),
        ("Comisión técnica", "Verificar fuentes, localizar brechas, preparar alternativas y documentar supuestos para deliberación."),
        ("Gobiernos locales y consejos municipales", "Aportar PMD, demandas, cartografía, información operativa y conocimiento de municipios y distritos municipales."),
        ("Representación social, económica y sectorial", "Contrastar evidencia, expresar perspectivas territoriales y acompañar el seguimiento."),
    ]
    add_table(doc, ["Instancia", "Función en la formulación"], roles, [2.15, 4.35], font_size=8.5)
    doc.add_heading("1.2 ¿Qué es este documento?", level=2)
    add_body_paragraph(
        doc,
        "Es una base técnica de trabajo para que el CDP formule el Plan Estratégico de Desarrollo Territorial correspondiente al nivel provincial. Ordena evidencia, plantea preguntas de validación y ofrece formatos para registrar acuerdos futuros.",
    )
    add_note_box(
        doc,
        "Alcance del documento base",
        "Organiza evidencia y formatos para que el CDP formule y concerte la vigencia, las prioridades, los compromisos presupuestarios y la cartera de proyectos.",
    )
    page_break(doc)

    doc.add_heading("1.3 Identidad y perfil territorial", level=2)
    general_rows = [
        ("Provincia", PROVINCE, "Clasificador geográfico"),
        ("Región de planificación", REGION, "Dashboard Territorial"),
        ("Código provincial", "14", "Clasificador geográfico"),
        ("Población", f"{facts['province']['population']:,} habitantes", "X Censo 2022, ONE"),
        ("Municipios", "Nagua, Cabrera, El Factor y Río San Juan", "División político-administrativa"),
        ("Hogares", f"{facts['households']['hogares_total']:,}", "X Censo 2022, ONE"),
        ("Población urbana / rural", f"{facts['urban_rural']['urbana']:,} / {facts['urban_rural']['rural']:,}", "X Censo 2022, ONE"),
    ]
    add_table(doc, ["Dato territorial", "Valor", "Fuente / año"], general_rows, [1.55, 2.9, 2.05], font_size=8.5)
    doc.add_heading("1.4 Antecedentes y alcance", level=2)
    add_body_paragraph(
        doc,
        "No se ha identificado un Plan Provincial publicado para María Trinidad Sánchez dentro del portal utilizado como línea base. La ausencia en el inventario no demuestra que no existan documentos, acuerdos o ejercicios anteriores; el CDP deberá completar esa verificación documental.",
    )
    add_body_paragraph(
        doc,
        "El Plan Provincial debe articular aportes de los municipios y sectores sin sustituir sus competencias ni copiar automáticamente los planes municipales. La escala provincial permite tratar interdependencias, servicios compartidos, movilidad, cuencas, redes de equipamientos y prioridades de inversión que superan un solo municipio.",
    )
    doc.add_heading("1.5 Información que debe completar el CDP", level=2)
    for item in [
        "Acta de conformación, miembros, sectores representados y mecanismo de sustitución.",
        "Inventario de planes, estudios, cartografía y acuerdos territoriales anteriores.",
        "Calendario de sesiones, metodología de validación y reglas para documentar disensos.",
        "Horizonte del plan, periodicidad de seguimiento y mecanismo de actualización.",
    ]:
        add_bullet(doc, item)
    page_break(doc)

    # Legal frame
    doc.add_heading("1.6 Marco jurídico e institucional", level=2)
    add_body_paragraph(
        doc,
        "El marco aplicable reconoce los consejos de desarrollo en tres niveles territoriales y les asigna participación en la formulación, priorización y seguimiento del desarrollo territorial.",
    )
    legal_rows = [
        (
            "Ley núm. 498-06",
            "Art. 14",
            "Constituye Consejos de Desarrollo en los niveles municipal, provincial y regional y establece su participación en los Planes Estratégicos de Desarrollo Territorial correspondientes.",
        ),
        (
            "Ley núm. 498-06",
            "Art. 15",
            "Asigna discusión y propuesta de estrategias, participación social, priorización de inversión, formulación de planes territoriales y seguimiento de su ejecución.",
        ),
        (
            "Ley núm. 498-06",
            "Arts. 30 y 36–38",
            "Vincula la planificación con la programación, evaluación y ciclo de la inversión pública.",
        ),
        (
            "Decreto núm. 493-07",
            "Arts. 4–6",
            "Desarrolla el carácter consultivo de los consejos y la consolidación ascendente de demandas y propuestas territoriales.",
        ),
        (
            "Decreto núm. 493-07",
            "Arts. 7–15",
            "Regula integración, coordinación, secretaría, reuniones, actas, comisión técnica y comunicación institucional.",
        ),
        (
            "Decreto núm. 493-07",
            "Arts. 57–61",
            "Articula la identificación y evaluación de iniciativas con el Sistema Nacional de Inversión Pública.",
        ),
    ]
    add_table(doc, ["Norma", "Artículos", "Aplicación a la formulación provincial"], legal_rows, [1.45, 1.05, 4.0], font_size=8)
    add_note_box(
        doc,
        "Nota institucional",
        "La Ley núm. 498-06 reconoce los Consejos de Desarrollo en los niveles municipal, provincial y regional y establece su participación en la formulación de los Planes Estratégicos de Desarrollo Territorial correspondientes. En este marco, el presente documento constituye una base técnica para la formulación del Plan Provincial, la consolidación de las demandas municipales y la definición de estrategias y prioridades de inversión. Su concertación, aprobación, seguimiento y actualización deberán ser acordados por el Consejo de Desarrollo Provincial y las autoridades competentes, de conformidad con la Ley núm. 498-06 y su Reglamento de Aplicación aprobado mediante el Decreto núm. 493-07.",
        fill="EEF5FA",
        title_color="2E74B5",
    )
    add_body_paragraph(
        doc,
        "Fuente: Ley núm. 498-06, arts. 14–15, 30 y 36–38; Decreto núm. 493-07, arts. 4–15 y 57–61.",
        source=True,
    )
    page_break(doc)

    # Methodology and diagnostic introduction
    doc.add_heading("1.7 Metodología de formulación", level=2)
    methodology = [
        ("1", "Organización", "Completar integración del CDP, reglas de trabajo, comisión técnica, cronograma y trazabilidad."),
        ("2", "Validación del diagnóstico", "Revisar fuentes, años, definiciones y distribución territorial; incorporar conocimiento local comprobable."),
        ("3", "Priorización concertada", "Acordar criterios de urgencia, equidad, escala provincial, competencia, impacto y madurez sin convertir indicadores automáticamente en acciones."),
        ("4", "Formulación estratégica", "Definir visión, objetivos, resultados, indicadores y cartera únicamente mediante acuerdos documentados del CDP."),
        ("5", "Articulación y seguimiento", "Relacionar acuerdos con inversión, presupuestos, instituciones, riesgos, calendario, reportes y revisión periódica."),
    ]
    add_table(doc, ["Paso", "Etapa", "Producto a acordar"], methodology, [0.55, 1.45, 4.5], font_size=8.5)
    doc.add_heading("2. Diagnóstico provincial", level=1)
    add_body_paragraph(
        doc,
        "Las nueve láminas siguientes presentan una línea base visual. Cada lámina conserva el año y la fuente, separa datos completos de períodos parciales y termina con preguntas para la deliberación. Una gráfica muestra patrones; no prueba por sí sola causalidad, calidad del servicio ni prioridad.",
    )
    plates_rows = [
        ("1", "Territorio y distribución municipal"),
        ("2", "Estructura demográfica y hogares"),
        ("3", "Condición de vida y servicios"),
        ("4", "Condiciones sociales"),
        ("5", "Seguridad vial y convivencia"),
        ("6", "Educación y economía formal"),
        ("7", "Equipamientos de salud y deporte"),
        ("8", "Inversión pública"),
        ("9", "Dinámica constructiva y demandas"),
    ]
    add_table(doc, ["Lámina", "Contenido"], plates_rows, [0.8, 5.7], font_size=8.5)
    add_note_box(
        doc,
        "Cómo leer las láminas",
        "Preguntar siempre: ¿qué mide?, ¿qué no mide?, ¿en qué territorio ocurre?, ¿qué fuente debe verificarlo?, ¿qué competencia institucional interviene? La respuesta debe registrarse antes de formular una prioridad.",
    )
    page_break(doc)

    # Plates
    for index, path in enumerate(plate_paths):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(path), width=Inches(6.5), height=Inches(7.8))
        if index != len(plate_paths) - 1:
            page_break(doc)
    page_break(doc)

    # Technical reading
    doc.add_heading("2.1 Lectura técnica y límites de la evidencia", level=1)
    readings = [
        (
            "Territorio y población",
            "Nagua concentra 52.0% de la población provincial; Cabrera, El Factor y Río San Juan reúnen el resto. El agregado no muestra por sí mismo tiempos de viaje, conectividad, dispersión rural, exposición a riesgos ni acceso real a servicios.",
            "Localizar comunidades, corredores, cuencas y equipamientos compartidos; verificar cómo cambian las necesidades entre zonas urbanas y rurales.",
        ),
        (
            "Servicios básicos",
            "El Censo registra 61.0% de hogares con agua del acueducto dentro de la vivienda, 78.7% con inodoro y 79.5% con recogida municipal de residuos. Las diferencias entre municipios son visibles.",
            "Comprobar continuidad, presión, potabilidad, frecuencia, rutas, disposición final, costo y población no servida; distinguir competencias municipales y sectoriales.",
        ),
        (
            "Condiciones sociales",
            "Los registros 2026 muestran 3.9% de hacinamiento extremo y 9.7% moderado; 13 centros INAIPI y 70.7% de asistencia registrada; 13,064 personas con discapacidad distribuidas por ICV.",
            "Revisar definiciones, evitar tratar presencias como personas únicas y localizar barreras específicas antes de definir población objetivo.",
        ),
        (
            "Seguridad y movilidad",
            "La tasa de homicidios fue 8.3 por 100 mil en 2024. La tasa de muertes de tránsito fue 32.6 por 100 mil en 2025, con variación anual considerable.",
            "Desagregar eventos por lugar, vía, horario, sexo, edad y circunstancia; contrastar infraestructura, control, educación y respuesta de emergencia.",
        ),
        (
            "Educación y economía",
            "El Anuario 2024 reporta 87.0% de promoción, 5.9% de abandono y 7.1% de reprobación en secundaria. El DEE registra 1,289 establecimientos y empleo estimado de 9,818 personas.",
            "Validar distritos educativos y evitar representar el DEE como totalidad de la economía, pues no cubre adecuadamente informalidad, autoempleo ni producción agropecuaria familiar.",
        ),
        (
            "Inversión y demandas",
            "El corte 2026 asocia 50 proyectos, RD$4.83 mil millones de presupuesto y 42.8% de ejecución registrada. El consolidado del CDP incluye diez demandas.",
            "Revisar localización, estado, duplicidades, contratos, madurez técnica, código SNIP, financiamiento y relación con brechas; no equiparar demanda, proyecto aprobado e inversión ejecutada.",
        ),
    ]
    for title, observation, verification in readings:
        doc.add_heading(title, level=2)
        add_body_paragraph(doc, "Dato observado. " + observation, bold_lead="Dato observado.")
        add_body_paragraph(doc, "Verificación requerida. " + verification, bold_lead="Verificación requerida.")
    page_break(doc)

    # Strengths and weaknesses
    doc.add_heading("3. Fortalezas y debilidades para validación", level=1)
    add_body_paragraph(
        doc,
        "La matriz convierte evidencia en preguntas de trabajo. Las categorías son provisionales: una cobertura alta puede ocultar problemas de calidad y una brecha estadística puede corresponder a distintas competencias o causas. El CDP debe confirmar, modificar o descartar cada fila.",
    )
    matrix_rows = [
        (
            "Fortaleza potencial",
            "Red territorial de equipamientos",
            "61 instalaciones deportivas; 39 de propiedad municipal; centros de salud presentes en los cuatro municipios.",
            "Puede existir una base física para coordinación intermunicipal, si el estado y la capacidad son adecuados.",
            "¿Qué instalaciones están operativas, accesibles y sirven a más de un municipio?",
        ),
        (
            "Fortaleza potencial",
            "Capacidad de inversión observada",
            "Ejecución registrada entre 86.4% y 98.8% en 2018–2025; 50 proyectos asociados en 2026.",
            "Existe flujo de inversión que podría articularse con acuerdos territoriales; las series requieren conciliación.",
            "¿Qué proyectos se alinean con brechas validadas y cuál es su estado real?",
        ),
        (
            "Fortaleza potencial",
            "Base institucional de demanda",
            "Diez demandas priorizadas y consolidadas por el CDP en 2026.",
            "Hay un punto de partida participativo documentado, sin que ello pruebe factibilidad o vigencia.",
            "¿Cuáles deben ratificarse, reformularse o agruparse y con qué criterio?",
        ),
        (
            "Debilidad / brecha potencial",
            "Agua y saneamiento",
            "61.0% con acueducto dentro; 78.7% con inodoro; diferencias municipales marcadas.",
            "El agregado sugiere brechas de acceso, pero no mide calidad, continuidad ni localización.",
            "¿Qué comunidades y sistemas explican la brecha y qué institución tiene competencia?",
        ),
        (
            "Debilidad / brecha potencial",
            "Residuos y drenaje",
            "79.5% declara recogida municipal; 16.9% quema residuos; demandas incluyen drenaje y saneamiento.",
            "Pueden existir brechas de cobertura y gestión ambiental; demanda y dato censal deben contrastarse.",
            "¿Qué rutas, frecuencias, puntos de disposición y zonas inundables deben verificarse?",
        ),
        (
            "Debilidad / brecha potencial",
            "Seguridad vial",
            "Tasa de muertes de tránsito de 32.6 por 100 mil en 2025.",
            "La serie señala un asunto para investigación, no una causa ni una solución predeterminada.",
            "¿Dónde, cuándo y en qué condiciones ocurren los eventos?",
        ),
        (
            "Debilidad / brecha potencial",
            "Inclusión y vivienda",
            "13.6% de hogares con hacinamiento extremo o moderado; 13,064 personas con discapacidad por ICV.",
            "La combinación puede indicar vulnerabilidades diferenciadas que requieren localización y validación.",
            "¿Qué hogares, barreras y servicios concentran la situación?",
        ),
        (
            "Debilidad / brecha potencial",
            "Información para decisiones",
            "Varios indicadores carecen de desagregación comunitaria o miden registros, no calidad ni personas únicas.",
            "La falta de detalle limita indicadores, metas y focalización confiables.",
            "¿Qué levantamientos mínimos debe completar la comisión técnica antes de formular acciones?",
        ),
    ]
    table = add_table(
        doc,
        ["Carácter", "Tema", "Evidencia", "Lectura provisional", "Pregunta al CDP"],
        matrix_rows,
        [1.0, 1.0, 1.6, 1.45, 1.45],
        font_size=7.2,
        alternating=False,
    )
    for row_index, row in enumerate(table.rows[1:], start=1):
        fill = "EEF7F4" if matrix_rows[row_index - 1][0].startswith("Fortaleza") else "FCF6E9"
        for cell in row.cells:
            set_cell_shading(cell, fill)
    add_body_paragraph(
        doc,
        "Resultado de la sesión: el CDP deberá marcar cada fila como confirmada, modificada, descartada o pendiente y anexar la evidencia local correspondiente.",
        small=True,
    )
    page_break(doc)

    # Concertation agenda
    doc.add_heading("4. Del diagnóstico a la concertación", level=1)
    doc.add_heading("4.1 Secuencia sugerida para las sesiones del CDP", level=2)
    session_rows = [
        ("1. Validar", "Confirmar fuente, año, definición, localización y responsable de aportar evidencia adicional."),
        ("2. Explicar", "Distinguir síntoma, causa posible, población afectada y competencias públicas o sociales."),
        ("3. Comparar", "Contrastar municipios, zonas urbanas/rurales, grupos de población e inversiones existentes."),
        ("4. Acordar", "Documentar situación priorizada, resultado esperado y criterio de selección; registrar disenso."),
        ("5. Formular", "Definir acción, indicador, línea base, meta, responsable, financiamiento, plazo y riesgos."),
    ]
    add_table(doc, ["Momento", "Pregunta de trabajo"], session_rows, [1.15, 5.35], font_size=8.5)
    doc.add_heading("4.2 Registro de validación del diagnóstico", level=2)
    validation_rows = [("\n\n", "\n\n", "\n\n", "\n\n", "\n\n") for _ in range(5)]
    validation_table = add_table(
        doc,
        ["Tema", "Dato validado", "Localización", "Fuente adicional", "Acuerdo / pendiente"],
        validation_rows,
        [1.15, 1.45, 1.15, 1.35, 1.4],
        font_size=8,
        alternating=False,
    )
    for row in validation_table.rows[1:]:
        for cell in row.cells:
            set_cell_shading(cell, "F3F6F6")
    doc.add_heading("4.3 Criterios de priorización a acordar", level=2)
    criteria_rows = [
        ("Urgencia y gravedad", "POR ACORDAR POR EL CDP", ""),
        ("Equidad territorial y poblacional", "POR ACORDAR POR EL CDP", ""),
        ("Escala o interdependencia provincial", "POR ACORDAR POR EL CDP", ""),
        ("Competencia y coordinación institucional", "POR ACORDAR POR EL CDP", ""),
        ("Madurez técnica y financiera", "POR ACORDAR POR EL CDP", ""),
        ("Sostenibilidad, riesgos y mantenimiento", "POR ACORDAR POR EL CDP", ""),
    ]
    add_table(doc, ["Criterio", "Definición acordada", "Peso / regla"], criteria_rows, [2.25, 3.25, 1.0], font_size=8)
    page_break(doc)

    # Vision and action intentionally blank
    doc.add_heading("5. Visión, objetivos y plan de acción", level=1)
    add_note_box(
        doc,
        "Sección reservada a decisión del CDP",
        "El diagnóstico permite formular preguntas, pero no autoriza a redactar una visión, seleccionar objetivos ni comprometer acciones. Los campos siguientes deben completarse mediante concertación y constar en actas.",
        fill="FCF6E9",
        title_color="B57A16",
    )
    doc.add_heading("5.1 Visión provincial", level=2)
    vision_rows = [
        ("Horizonte temporal", "POR ACORDAR POR EL CDP"),
        ("Texto de la visión", "POR ACORDAR POR EL CDP\n\n"),
        ("Acta / fecha de concertación", "POR ACORDAR POR EL CDP"),
        ("Participantes y constancia de disenso", "POR ACORDAR POR EL CDP"),
    ]
    add_table(doc, ["Campo", "Acuerdo"], vision_rows, [2.0, 4.5], font_size=9, alternating=False)
    doc.add_heading("5.2 Objetivos y resultados", level=2)
    objectives_rows = [("\n\n", "\n\n", "\n\n", "\n\n", "\n\n") for _ in range(4)]
    objectives_table = add_table(
        doc,
        ["Situación priorizada", "Objetivo", "Resultado", "Indicador / línea base", "Meta"],
        objectives_rows,
        [1.35, 1.25, 1.25, 1.65, 1.0],
        font_size=7.8,
        alternating=False,
    )
    for row in objectives_table.rows[1:]:
        for cell in row.cells:
            set_cell_shading(cell, "F3F6F6")
    page_break(doc)

    doc.add_heading("5.3 Ficha para formular una acción o proyecto", level=2)
    action_rows = [
        ("Situación priorizada y evidencia", "POR ACORDAR POR EL CDP"),
        ("Resultado esperado", "POR ACORDAR POR EL CDP"),
        ("Acción / proyecto", "POR ACORDAR POR EL CDP"),
        ("Alcance territorial y población", "POR ACORDAR POR EL CDP"),
        ("Indicador, línea base y meta", "POR ACORDAR POR EL CDP"),
        ("Institución responsable y aliadas", "POR ACORDAR POR EL CDP"),
        ("Presupuesto, fuente y código SNIP", "POR ACORDAR POR EL CDP"),
        ("Plazo, hitos y mantenimiento", "POR ACORDAR POR EL CDP"),
        ("Riesgos, supuestos y verificación", "POR ACORDAR POR EL CDP"),
        ("Acta / fecha del acuerdo", "POR ACORDAR POR EL CDP"),
    ]
    add_table(doc, ["Campo", "Definición concertada"], action_rows, [2.2, 4.3], font_size=8.5, alternating=False)
    doc.add_heading("5.4 Matriz de seguimiento", level=2)
    follow_rows = [("\n\n", "\n\n", "\n\n", "\n\n", "\n\n") for _ in range(4)]
    follow_table = add_table(
        doc,
        ["Resultado / acción", "Indicador", "Meta / plazo", "Responsable", "Avance y evidencia"],
        follow_rows,
        [1.45, 1.2, 1.2, 1.25, 1.4],
        font_size=7.8,
        alternating=False,
    )
    for row in follow_table.rows[1:]:
        for cell in row.cells:
            set_cell_shading(cell, "F3F6F6")
    page_break(doc)

    # Demands appendix
    doc.add_heading("Anexo A. Demandas provinciales consolidadas", level=1)
    add_body_paragraph(
        doc,
        "El consolidado 003 registra diez demandas de María Trinidad Sánchez priorizadas en 2026. Se incluyen como evidencia de participación y como insumo para contrastar con el diagnóstico. El CDP deberá revisar vigencia, alcance, duplicidades, institución responsable, código SNIP, madurez técnica y financiamiento.",
    )
    demand_rows = []
    for item in facts["demands"]:
        title = re.sub(r"^2027[_\s-]*", "", item["demand"], flags=re.I)
        demand_rows.append(
            (
                item["number"],
                title,
                item["municipality"].title() if item["municipality"] else "Provincial / sin municipio",
                shorten_institution(item["institution"]),
                item["snip"] or "Pendiente",
            )
        )
    add_table(
        doc,
        ["Núm.", "Demanda priorizada (texto resumido)", "Municipio", "Institución", "SNIP"],
        demand_rows,
        [0.45, 3.1, 1.05, 1.05, 0.85],
        font_size=7.2,
        alternating=True,
    )
    add_note_box(
        doc,
        "Decisión pendiente",
        "Una demanda consolidada no equivale automáticamente a una acción del Plan Provincial. La incorporación debe quedar documentada mediante criterios de priorización, diagnóstico validado y ficha técnica.",
        fill="EEF5FA",
        title_color="2E74B5",
    )
    page_break(doc)

    # Sources
    doc.add_heading("Fuentes y trazabilidad", level=1)
    source_rows = [
        ("ONE-2022", "X Censo Nacional de Población y Vivienda", "2022", "Población, hogares, servicios y educación", "Alta"),
        ("DASH-DIAG", "Dashboard de Diagnóstico Territorial y datasets derivados", "2022 / 2024", "Comparaciones provincial/municipal, economía y equipamientos", "Alta / media según indicador"),
        ("DASH-PROV", "Dashboard Territorial", "2001–2026", "Seguridad, condiciones sociales, equipamientos, inversión, vías y permisos", "Según ficha de fuente"),
        ("INV-2026", "Inversión Pública Territorial", "2026", "50 proyectos, presupuesto y ejecución al 31-07-2026", "Corte administrativo"),
        ("DEM-2026", "Demandas Provinciales · consolidado 003", "2026", "Diez demandas de María Trinidad Sánchez", "Registro del CDP"),
        ("LEY-498", "Ley núm. 498-06 de Planificación e Inversión Pública", "2006", "Arts. 14–15, 30 y 36–38", "Normativa oficial"),
        ("DEC-493", "Decreto núm. 493-07, Reglamento de Aplicación", "2007", "Arts. 4–15 y 57–61", "Normativa oficial"),
    ]
    add_table(doc, ["ID", "Fuente", "Año", "Cobertura", "Confianza / límite"], source_rows, [0.75, 2.3, 0.75, 1.65, 1.05], font_size=7.5)
    doc.add_heading("Enlaces de consulta", level=2)
    links = [
        ("Dashboard Territorial", "https://prodecare.net/DDPT/Dashboard-Territorial/"),
        ("Inversión Pública Territorial", "https://prodecare.net/DDPT/InversionPublicaTerritorial/"),
        ("Demandas Provinciales", "https://prodecare.net/DDPT/DemandasProvinciales/"),
        ("Planificación Municipal", "https://prodecare.net/DDPT/PlanificacionMunicipal/"),
        ("Ley núm. 498-06", "https://mepyd.gob.do/wp-content/uploads/drive/DIGEDES/Monitoreo%20y%20Evaluaci%C3%B3n/Publicaciones/Normativa/Ley-498-06%20Planificaci%C3%B3n%20e%20Inversi%C3%B3n%20P%C3%BAblica.pdf"),
        ("Decreto núm. 493-07", "https://mepyd.gob.do/wp-content/uploads/drive/DIGEDES/Monitoreo%20y%20Evaluaci%C3%B3n/Publicaciones/Normativa/Decreto-493-07%20Reglamento%20aplicaci%C3%B3n%20ley%20Planificaci%C3%B3n%20e%20Inversi%C3%B3n%20P%C3%BAblica.pdf"),
    ]
    for label, url in links:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(label + ": ")
        set_run_font(run, size=10, color="203740")
        add_hyperlink(paragraph, "Abrir fuente", url)
    add_note_box(
        doc,
        "Control de versión",
        "Documento base elaborado con información disponible al 1 de agosto de 2026. Antes de una sesión formal, la comisión técnica deberá revisar cambios de fuente, períodos parciales y actualizaciones institucionales.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


def collect_facts():
    TEMP.mkdir(parents=True, exist_ok=True)
    ensure_download(DASHBOARD_URL, DASHBOARD_PATH)
    ensure_download(DEMANDS_URL, DEMANDS_PATH)
    ensure_download(REFERENCE_URL, REFERENCE_PATH)
    dashboard = load_json(DASHBOARD_PATH)
    province = next(item for item in dashboard["provinces"] if item["name"] == PROVINCE)
    municipalities = [item for item in dashboard["municipalities"] if item.get("province") == PROVINCE]
    portal_data = load_json(PROJECT / "src" / "data" / "provinces.json")
    investment_2026 = next(item["investment"] for item in portal_data["provinces"] if item["name"] == PROVINCE)
    return {
        "province": province,
        "municipalities": municipalities,
        "investment_2026": investment_2026,
        "demands": parse_demands(DEMANDS_PATH),
        "geojson": load_json(PROJECT / "public" / "data" / "provinces.geojson"),
        "urban_rural": find_record(SOURCE_DATA / "poblacion_urbana_rural_provincia.json"),
        "households": find_record(SOURCE_DATA / "hogares_resumen_provincia.json"),
        "pyramid": find_record(SOURCE_DATA / "pyramids_provincia.json"),
        "condition": find_record(SOURCE_DATA / "condicion_vida_provincia.json"),
        "municipal_condition": [item for item in load_json(SOURCE_DATA / "condicion_vida.json") if item.get("provincia") == PROVINCE],
        "tic": find_record(SOURCE_DATA / "tic_provincia.json"),
        "education_level": find_record(SOURCE_DATA / "educacion_nivel_provincia.json"),
        "education": find_record(SOURCE_DATA / "educacion_provincia.json"),
        "economy": find_record(SOURCE_DATA / "economia_empleo_provincia.json"),
        "health": find_record(SOURCE_DATA / "salud_establecimientos_provincia.json"),
    }


def validate_facts(facts):
    assert facts["province"]["population"] == 156633
    assert len(facts["municipalities"]) == 4
    assert sum(item["population"] for item in facts["municipalities"]) == facts["province"]["population"]
    assert len(facts["demands"]) == 10
    assert facts["investment_2026"]["projectCount"] == 50
    assert abs(facts["investment_2026"]["executionPct"] - 42.8331359556) < 0.01


def main():
    facts = collect_facts()
    validate_facts(facts)
    plates = build_plates(facts)
    output = build_document(facts, plates)
    print(json.dumps({"output": str(output), "plates": [str(path) for path in plates], "demands": len(facts["demands"])}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
