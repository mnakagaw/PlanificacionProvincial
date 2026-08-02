from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from PIL import Image

import build_mts_plan as builder


PORTFOLIO_HEADERS = [
    "N.º",
    "Identificación",
    "Proyecto / institución / territorio",
    "Estado / período / avances",
    "Costos y ejecución",
]


def table_headers(table):
    return [cell.text.strip() for cell in table.rows[0].cells]


def require(condition, message):
    if not condition:
        raise AssertionError(message)


def main():
    parser = argparse.ArgumentParser(description="Valida los 32 documentos provinciales generados.")
    parser.add_argument("directory", type=Path)
    args = parser.parse_args()
    directory = args.directory.resolve()

    portal = builder.load_json(builder.PROJECT / "src" / "data" / "provinces.json")
    investment = builder.load_json(builder.INVESTMENT_PATH)
    age_2010 = builder.load_json(builder.SOURCE_DATA / "edad_sexo_2010_provincia.json")
    records = sorted(portal["provinces"], key=lambda item: int(builder.PROVINCE_CODES[item["name"]]))
    manifest = builder.load_json(directory / "provincial-documents.json")
    manifest_by_key = {item["provinceKey"]: item for item in manifest["documents"]}

    require(len(records) == 32, f"Se esperaban 32 provincias; hay {len(records)}")
    require(len(manifest_by_key) == 32, f"El manifiesto contiene {len(manifest_by_key)} documentos")

    project_total = 0
    page_minimum_total = 0
    validations = []
    for record in records:
        file_name = builder.set_province_context(record)
        projects = builder.select_investment_projects(investment)
        doc_path = directory / file_name
        require(doc_path.exists(), f"Falta {doc_path.name}")
        document = Document(doc_path)
        portfolios = [table for table in document.tables if table_headers(table) == PORTFOLIO_HEADERS]
        require(len(portfolios) == 1, f"{builder.PROVINCE}: tablas de cartera={len(portfolios)}")
        portfolio = portfolios[0]
        require(len(portfolio.rows) == len(projects) + 1, f"{builder.PROVINCE}: filas de cartera incorrectas")

        all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        require("B.2 Localización y calendario" not in all_text, f"{builder.PROVINCE}: conserva B.2")
        require("B.3 Costos, ejecución" not in all_text, f"{builder.PROVINCE}: conserva B.3")

        relationship_targets = {
            str(relationship.target_ref)
            for relationship in document.part.rels.values()
            if getattr(relationship, "target_ref", None)
        }
        for row, project in zip(portfolio.rows[1:], projects):
            identification = row.cells[1].text
            detail = row.cells[2].text
            status = row.cells[3].text
            finance = row.cells[4].text
            budget = float(project.get("budget") or 0)
            executed = float(project.get("executed") or 0)
            execution = builder.format_percent(executed / budget * 100) if budget else "No calculable"

            require(f"SNIP {project.get('code') or 'Sin dato'}" in identification, f"{builder.PROVINCE}: SNIP ausente")
            require(f"Ficha/Mapa {project.get('mapProjectId') or 'Sin dato'}" in identification, f"{builder.PROVINCE}: ficha ausente")
            require((project.get("name") or "Sin dato") in detail, f"{builder.PROVINCE}: nombre ausente")
            require(f"Localización codificada: {builder.project_location_rows(project)}" in detail, f"{builder.PROVINCE}: localización ausente")
            require(f"Estado: {project.get('state') or 'Sin dato'}" in status, f"{builder.PROVINCE}: estado ausente")
            require(builder.format_percent(project.get("physicalProgress"), ratio=True) in status, f"{builder.PROVINCE}: avance físico ausente")
            require(builder.format_percent(project.get("financialProgress"), ratio=True) in status, f"{builder.PROVINCE}: avance financiero ausente")
            require(builder.format_rd(project.get("projectCost")) in finance, f"{builder.PROVINCE}: costo ausente")
            require(builder.format_rd(budget) in finance, f"{builder.PROVINCE}: presupuesto ausente")
            require(builder.format_rd(executed) in finance, f"{builder.PROVINCE}: ejecutado ausente")
            require(execution in finance, f"{builder.PROVINCE}: tasa ausente")
            if project.get("projectUrl"):
                require(project["projectUrl"] in relationship_targets, f"{builder.PROVINCE}: enlace de proyecto ausente")
            if project.get("contractUrl"):
                require(project["contractUrl"] in relationship_targets, f"{builder.PROVINCE}: enlace de contrato ausente")

        province_age = [
            item
            for item in age_2010
            if builder.normalize(item.get("provincia")) == builder.normalize(builder.PROVINCE)
        ]
        require(len(province_age) == 21, f"{builder.PROVINCE}: grupos 2010={len(province_age)}")
        pyramid_path = builder.CHARTS / "lamina_02_demografia.png"
        require(pyramid_path.exists(), f"{builder.PROVINCE}: falta la lámina demográfica")
        with Image.open(pyramid_path).convert("RGB") as image:
            colors = set(image.getdata())
        for color in (builder.rgb(builder.COLORS["blue"]), builder.rgb(builder.COLORS["red"]), builder.rgb("#59666B"), builder.rgb("#AEB7BA")):
            require(color in colors, f"{builder.PROVINCE}: falta color de pirámide {color}")

        manifest_record = manifest_by_key[record["key"]]
        require(manifest_record["fileName"] == file_name, f"{builder.PROVINCE}: nombre de manifiesto incorrecto")
        require(manifest_record["size"] == doc_path.stat().st_size, f"{builder.PROVINCE}: tamaño de manifiesto incorrecto")
        project_total += len(projects)
        page_minimum_total += int(manifest_record["pagesExpectedMinimum"])
        validations.append({"province": builder.PROVINCE, "projects": len(projects), "bytes": doc_path.stat().st_size})

    print(
        json.dumps(
            {
                "documents": len(validations),
                "projects_across_documents": project_total,
                "minimum_pages_declared": page_minimum_total,
                "smallest_document": min(item["bytes"] for item in validations),
                "largest_document": max(item["bytes"] for item in validations),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
